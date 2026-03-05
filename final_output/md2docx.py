#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown to Word (docx) converter with support for tables, headings, and basic formatting.
"""

import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def parse_markdown_to_docx(md_file_path, docx_file_path):
    """Convert markdown file to Word document."""

    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

    # Split content into lines
    lines = content.split('\n')

    i = 0
    in_table = False
    table = None

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Handle headings
        if line.startswith('# '):
            add_heading(doc, line[2:], 1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], 2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], 3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], 4)
        # Handle table
        elif line.startswith('|'):
            if not in_table:
                in_table = True
                # Parse table header
                header_line = line
                i += 1
                if i < len(lines) and lines[i].startswith('|'):
                    separator_line = lines[i]
                    # Get column count from separator
                    cols = len([x for x in header_line.split('|') if x.strip()])
                    # Create table
                    table = doc.add_table(rows=0, cols=cols)
                    table.style = 'Table Grid'
                    # Add header row
                    header_cells = table.add_row().cells
                    header_parts = [x.strip() for x in header_line.split('|') if x.strip()]
                    for j, part in enumerate(header_parts):
                        if j < len(header_cells):
                            header_cells[j].text = part

                    # Skip separator
                    i += 1

                # Parse data rows
                while i < len(lines) and lines[i].startswith('|'):
                    data_line = lines[i]
                    data_parts = [x.strip() for x in data_line.split('|') if x.strip()]
                    if len(data_parts) > 0:  # Skip separator lines
                        row_cells = table.add_row().cells
                        for j, part in enumerate(data_parts):
                            if j < len(row_cells):
                                row_cells[j].text = part
                    i += 1
                continue  # Skip the i+=1 at end
            else:
                # Continue table parsing
                data_parts = [x.strip() for x in line.split('|') if x.strip()]
                if len(data_parts) > 0:
                    row_cells = table.add_row().cells
                    for j, part in enumerate(data_parts):
                        if j < len(row_cells):
                            row_cells[j].text = part
        else:
            in_table = False
            # Handle regular paragraph with inline formatting
            # Check for image
            if line.startswith('![图'):
                # Image line - add as paragraph with image description
                # Extract image path from ![alt](path)
                img_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
                if img_match:
                    alt_text = img_match.group(1)
                    img_path = img_match.group(2)
                    p = doc.add_paragraph()
                    p.add_run(f'[图片: {alt_text}]').italic = True
            else:
                # Regular paragraph
                add_paragraph_with_formatting(doc, line)

        i += 1

    # Save document
    doc.save(docx_file_path)
    print(f"Converted: {md_file_path} -> {docx_file_path}")

def add_heading(doc, text, level):
    """Add heading with proper styling."""
    heading = doc.add_heading(text, level=level)
    # Set heading font
    for run in heading.runs:
        run.font.name = '黑体'
        run.font.size = Pt(14) if level <= 2 else Pt(12)
        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph_with_formatting(doc, text):
    """Add paragraph with basic inline formatting (bold, italic, math)."""
    p = doc.add_paragraph()

    # Process text for inline formatting
    # Handle bold: **text**
    # Handle italic: *text*
    # Handle math: $...$ or $$...$$

    remaining = text

    while remaining:
        # Find next special marker
        bold_match = re.search(r'\*\*([^*]+)\*\*', remaining)
        italic_match = re.search(r'\*([^*]+)\*', remaining)
        math_match = re.search(r'\$([^\$]+)\$', remaining)
        equation_match = re.search(r'\$\$([^$]+)\$\$', remaining)

        # Find earliest match
        matches = []
        if bold_match:
            matches.append(('bold', bold_match.start(), bold_match.end(), bold_match.group(1)))
        if italic_match:
            matches.append(('italic', italic_match.start(), italic_match.end(), italic_match.group(1)))
        if math_match:
            matches.append(('math', math_match.start(), math_match.end(), math_match.group(1)))
        if equation_match:
            matches.append(('equation', equation_match.start(), equation_match.end(), equation_match.group(1)))

        if not matches:
            # No more special formatting, add remaining text
            if remaining.strip():
                p.add_run(remaining)
            break

        # Sort by position
        matches.sort(key=lambda x: x[1])
        earliest = matches[0]

        # Add text before the match
        if earliest[1] > 0:
            p.add_run(remaining[:earliest[1]])

        # Add formatted text
        run = p.add_run(earliest[3])
        if earliest[0] == 'bold':
            run.bold = True
        elif earliest[0] == 'italic':
            run.italic = True
        elif earliest[0] in ('math', 'equation'):
            # Keep math as-is for now (would need equation editor for full support)
            run.font.name = 'Times New Roman'

        remaining = remaining[earliest[2]:]

    # Set paragraph alignment
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

if __name__ == '__main__':
    import glob

    # Convert 5.1 and 5.2 files
    base_dir = r'D:\桌面\lunwen\thesis\final_output'
    output_dir = r'D:\桌面\lunwen\thesis\final_output\docs'

    files = [
        ('第5章_5.1_宽带LFMCW诊断系统设计与时间分辨率测试_final.md',
         '第5章_5.1_宽带LFMCW诊断系统设计与时间分辨率测试_new.docx'),
        ('第5章_5.2_微波带通滤波器的色散物理等效机理_final.md',
         '第5章_5.2_微波带通滤波器的色散物理等效机理_new.docx'),
    ]

    for md_name, docx_name in files:
        md_path = os.path.join(base_dir, md_name)
        docx_path = os.path.join(output_dir, docx_name)

        if os.path.exists(md_path):
            parse_markdown_to_docx(md_path, docx_path)
        else:
            print(f"File not found: {md_path}")
